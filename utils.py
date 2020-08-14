####################################
# Author: Jeremy (Meng-Chieh) Lee  #
# Date	: 2020/06/02               #
####################################

import os
import numpy as np
from math import ceil
import pandas as pd
from collections import defaultdict
import progressbar

import string
from nltk.corpus import stopwords

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

WCI = {-1: WD_COLOR_INDEX.RED,
		0: WD_COLOR_INDEX.YELLOW, \
		1: WD_COLOR_INDEX.BRIGHT_GREEN, \
		2: WD_COLOR_INDEX.GRAY_25, \
		3: WD_COLOR_INDEX.TEAL}

def set_global_voc_cost(c):
	global GOLBAL_VOC_COST
	GOLBAL_VOC_COST = c

def log_star(x):
	"""
	Universal code length

	"""
	return 2 * ceil(np.log2(x)) + 1 if x != 0 else 0

def word_cost():
	return GOLBAL_VOC_COST

def sequence_cost(seq):
	"""
	Output encoding cost for a given sequence

	"""
	return log_star(len(seq)) + len(seq) * word_cost()

def str_prep(s):
	s = s.translate(str.maketrans('', '', string.punctuation)).split(' ')
	s = np.array([ss.lower() for ss in s if len(ss) != 0])
	return s

def read_data(path):
	df = pd.read_csv(path)
	lsh_label = df['LSH label'].unique()
	data = defaultdict(dict)

	voc = set()
	for label in progressbar.progressbar(lsh_label):
		for id, text in df[df['LSH label'] == label][['id', 'text']].values:
			try:
				text = str_prep(text)
				for t in text:
					voc.add(t)
			except:
				continue
			if len(text) != 0:
				data[label][id] = text

	gvc = ceil(np.log2(len(voc)))
	# set_global_voc_cost(gvc)
	return data, gvc

def output_word(temp, cond, word_path):
	"""
	Output highlight content with office word document

	"""
	### Initialize document
	doc = Document()
	proc = doc.add_paragraph()
	for s, c in zip(['Slot', 'Matched', 'Substitution', 'Deletion', 'Insertion'], WCI.values()):
		font = proc.add_run(s).font
		font.highlight_color = c
		proc.add_run(' ')

	### Template content
	proc = doc.add_paragraph()
	proc.add_run('Template: \n')
	proc.add_run(temp.seq())
	proc.add_run('\n\n-----------------------------------------------------------------\n')

	### Iterate all aligned sequences
	for cs in cond:
		proc = doc.add_paragraph()
		for c, s in cs:
			font = proc.add_run(s).font
			font.highlight_color = WCI[c]
			proc.add_run(' ')

	doc.save(word_path)

def output_results(temp_arr, cond_arr, output_path, html_name='graph.html', word_name='text.docx'):
	"""
	Output template results

	"""
	if not os.path.exists(output_path):
		os.makedirs(output_path)

	### Iterate all templates
	for idx, (temp, cond) in enumerate(zip(temp_arr, cond_arr)):
		temp_path = os.path.join(output_path, 'template_' + str(idx + 1))
		if not os.path.exists(temp_path):
			os.makedirs(temp_path)

		### Output html
		temp.htmlOutput(open(os.path.join(temp_path, html_name), 'w'))

		### Output word document
		output_word(temp, cond, os.path.join(temp_path, word_name))
